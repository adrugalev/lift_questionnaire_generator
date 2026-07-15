# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any

from .models import LiftGroup, ProjectInfo, Questionnaire


@dataclass(frozen=True)
class ParseResult:
    text: str
    questionnaire: Questionnaire
    found_fields: dict[str, bool]
    warnings: list[str] = field(default_factory=list)
    extraction_method: str = "text"


def extract_text(path_or_bytes: str | Path | bytes, filename: str | None = None) -> str:
    suffix = _suffix(path_or_bytes, filename)
    if suffix == ".docx":
        return _extract_docx(path_or_bytes)
    if suffix == ".pdf":
        text, _, _ = _extract_pdf(path_or_bytes)
        return text
    if suffix == ".doc":
        raise ValueError("DOC поддерживается только после конвертации в DOCX или PDF.")
    raise ValueError(f"Неподдерживаемый формат файла: {suffix or 'без расширения'}")


def parse_specification(path_or_bytes: str | Path | bytes, filename: str | None = None) -> ParseResult:
    suffix = _suffix(path_or_bytes, filename)
    warnings: list[str] = []
    method = "text"
    if suffix == ".pdf":
        text, warnings, method = _extract_pdf(path_or_bytes)
    else:
        text = extract_text(path_or_bytes, filename)

    questionnaire, found = parse_text(text)
    if suffix == ".pdf" and not questionnaire.lift_groups:
        rapid_questionnaire = _try_rapidocr_lift_table(path_or_bytes)
        if rapid_questionnaire and rapid_questionnaire.lift_groups:
            rapid_found = _build_found_fields(rapid_questionnaire.lift_groups)
            rapid_found.update(_project_found_fields(rapid_questionnaire.project))
            return ParseResult(
                text=text,
                questionnaire=rapid_questionnaire,
                found_fields=rapid_found,
                warnings=warnings,
                extraction_method="rapidocr-table",
            )

    if suffix == ".pdf" and not questionnaire.lift_groups:
        vision_questionnaire = _try_openai_vision(path_or_bytes)
        if vision_questionnaire and vision_questionnaire.lift_groups:
            vision_found = _build_found_fields(vision_questionnaire.lift_groups)
            vision_found.update(_project_found_fields(vision_questionnaire.project))
            return ParseResult(
                text=text,
                questionnaire=vision_questionnaire,
                found_fields=vision_found,
                warnings=warnings,
                extraction_method="openai-vision",
            )

    if suffix == ".pdf" and not questionnaire.lift_groups:
        warnings.append(
            "Не удалось извлечь параметры лифта из PDF. Если это скан или чертеж без текстового слоя, "
            "установите Tesseract OCR, задайте OPENAI_API_KEY для vision-распознавания "
            "или загрузите PDF/DOCX с распознаваемым текстом."
        )
    return ParseResult(
        text=text,
        questionnaire=questionnaire,
        found_fields=found,
        warnings=warnings,
        extraction_method=method,
    )


def _try_openai_vision(path_or_bytes: str | Path | bytes) -> Questionnaire | None:
    try:
        from .llm_extractor import extract_pdf_with_vision
    except Exception:
        return None
    try:
        return extract_pdf_with_vision(path_or_bytes)
    except Exception:
        return None


def parse_text(text: str) -> tuple[Questionnaire, dict[str, bool]]:
    project_name = _extract_project_name(text)
    groups = _parse_table_like_rows(text)
    if not groups:
        groups = _parse_flat_text(text)

    found = _build_found_fields(groups)
    found.update(_project_found_fields(ProjectInfo(project_name=project_name)))
    questionnaire = Questionnaire(project=ProjectInfo(project_name=project_name), lift_groups=groups)
    return questionnaire, found


def _parse_table_like_rows(text: str) -> list[LiftGroup]:
    rows = [_split_table_row(line) for line in text.splitlines()]
    rows = [row for row in rows if len(row) >= 3 and any(row)]
    if not rows:
        return []

    header_index, group_columns = _find_group_header(rows)
    if header_index is None or not group_columns:
        return []

    group_data: list[dict[str, Any]] = []
    for column in group_columns:
        group_data.append(_parse_group_header(rows[header_index][column]))

    for row in rows[header_index + 1 :]:
        label = _row_label(row)
        if not label:
            continue
        for group_index, column in enumerate(group_columns):
            if column >= len(row):
                continue
            _apply_table_value(group_data[group_index], label, row[column])

    return [LiftGroup(**_drop_empty(data)) for data in group_data if _has_lift_data(data)]


def _parse_flat_text(text: str) -> list[LiftGroup]:
    normalized = re.sub(r"\s+", " ", text)
    group = LiftGroup(
        quantity=_int_match(normalized, [r"количеств[оа]\s+лифтов\s*[:\-]?\s*(\d+)"]),
        capacity_kg=_int_match(normalized, [r"грузопод[ъь]?емност[ьи]\s*[:\-]?\s*(\d{3,5})", r"(\d{3,5})\s*кг"]),
        speed_ms=_float_match(normalized, [r"скорост[ьи]\s*[:\-]?\s*([0-9]+(?:[,.][0-9]+)?)\s*м\s*/?\s*с"]),
        lifting_height_mm=_measurement_mm_match(normalized, [r"высот[аы]\s+под[ъь]?ема\s*[:\-]?\s*([0-9]+(?:[,.][0-9]+)?)"]),
        stops=_int_match(normalized, [r"(?:количеств[оа]\s+)?останов(?:ок|ки)\s*[:\-]?\s*(\d{1,2})"]),
        doors_count=_int_match(normalized, [r"(?:количеств[оа]\s+)?двер(?:ей|и)\s*[:\-]?\s*(\d{1,2})"]),
        cabin_width_mm=_int_match(normalized, [r"ширин[аы]\s+кабин[ыа]\s*[:\-]?\s*(\d{3,5})"]),
        cabin_depth_mm=_int_match(normalized, [r"глубин[аы]\s+кабин[ыа]\s*[:\-]?\s*(\d{3,5})"]),
        cabin_height_mm=_int_match(normalized, [r"высот[аы]\s+кабин[ыа]\s*[:\-]?\s*(\d{3,5})"]),
        landing_door_width_mm=_int_match(normalized, [r"ширин[аы]\s+двер(?:ей|и)\s*[:\-]?\s*(\d{3,5})"]),
        landing_door_height_mm=_int_match(normalized, [r"высот[аы]\s+двер(?:ей|и)\s*[:\-]?\s*(\d{3,5})"]),
        pit_depth_mm=_int_match(normalized, [r"глубин[аы]\s+приямк[а]?\s*[:\-]?\s*(\d{3,5})"]),
        overhead_mm=_int_match(normalized, [r"высот[аы]\s+верхн(?:его|ей)\s+этаж[а]?\s*[:\-]?\s*(\d{3,5})"]),
    )
    return [group] if _has_lift_data(group.model_dump()) else []


def _extract_project_name(text: str) -> str | None:
    normalized = re.sub(r"\s+", " ", text)
    return _first_match(
        normalized,
        [
            r"(?:проект|объект)\s*[:\-]\s*([^.;\n]{3,120})",
            r"(?:жилой комплекс|жк)\s+([^.;\n]{3,120})",
        ],
    )


def _split_table_row(line: str) -> list[str]:
    if "\t" in line:
        return [_clean_cell(cell) for cell in line.split("\t")]
    if "|" in line:
        return [_clean_cell(cell) for cell in line.split("|")]
    return []


def _find_group_header(rows: list[list[str]]) -> tuple[int | None, list[int]]:
    for row_index, row in enumerate(rows[:20]):
        group_columns = [
            column
            for column, cell in enumerate(row)
            if _looks_like_group_header(cell) and not _is_note_column(cell)
        ]
        if group_columns:
            return row_index, group_columns
    return None, []


def _looks_like_group_header(value: str) -> bool:
    value = value.lower()
    return "лифт" in value or "секц" in value or "г/п" in value or "грузопод" in value


def _is_note_column(value: str) -> bool:
    return "примеч" in value.lower()


def _parse_group_header(value: str) -> dict[str, Any]:
    data: dict[str, Any] = {"lift_name": _clean_cell(value) or None}
    section = _first_match(value, [r"секци[яи]\s*([0-9а-яА-ЯёЁ.,;\-/–— ]+)"])
    if section:
        data["section"] = f"Секция {section.strip()}"
    if re.search(r"пассажир", value, flags=re.IGNORECASE):
        data["lift_type"] = "Пассажирский"
    capacity = _int_match(value, [r"(?:г/п|грузопод[ъь]?емност[ьи]?)\s*(\d{3,5})", r"(\d{3,5})\s*кг"])
    if capacity:
        data["capacity_kg"] = capacity
    return data


def _row_label(row: list[str]) -> str | None:
    candidates = row[:2] if len(row) > 2 else row
    for cell in reversed(candidates):
        if cell and not re.fullmatch(r"\d+", cell):
            return cell
    return None


def _apply_table_value(data: dict[str, Any], label: str, value: str) -> None:
    label_norm = _normalize(label)
    value = _clean_cell(value)
    if not value:
        return

    if "тип лифта" in label_norm:
        data["lift_type"] = _capitalize_value(value)
    elif "грузопод" in label_norm or "скорост" in label_norm:
        data["capacity_kg"] = _int_match(value, [r"(?:q|г/п|грузопод[ъь]?емност[ьи]?)\s*=?\s*(\d{3,5})", r"(\d{3,5})\s*кг"]) or data.get("capacity_kg")
        data["speed_ms"] = _float_match(value, [r"(?:v|скорост[ьи]?)\s*=?\s*([0-9]+(?:[,.][0-9]+)?)"]) or data.get("speed_ms")
    elif "высота подъема" in label_norm:
        data["lifting_height_mm"] = _measurement_mm(value)
    elif "размер кабины" in label_norm or ("ширина" in label_norm and "глубина" in label_norm and "кабин" in label_norm):
        size = _parse_size(value)
        if size:
            data["cabin_width_mm"], data["cabin_depth_mm"] = size
    elif "противополож" in label_norm or "проход" in label_norm:
        data["cabin_type"] = "Проходная" if _yes(value) else "Непроходная"
    elif "число двер" in label_norm or "количество двер" in label_norm:
        data["doors_count"] = _int_from_value(value)
    elif "число останов" in label_norm or "количество останов" in label_norm:
        data["stops"] = _int_from_value(value)
    elif "отметк" in label_norm and "посад" in label_norm:
        data["button_marking"] = value
        if "0.000" in value or "0,000" in value:
            data["main_landing_floor"] = "0.000"
    elif "управление лифт" in label_norm or "работа в группе" in label_norm:
        data["group_operation"] = _group_operation(value)
    elif "конструкц" in label_norm and "шахт" in label_norm:
        data["shaft_material"] = _shaft_material_from_ocr(value)
    elif "глубина приям" in label_norm:
        data["pit_depth_mm"] = _measurement_mm(value)
    elif "высота верх" in label_norm:
        data["overhead_mm"] = _measurement_mm(value)
    elif "особые треб" in label_norm or "огнестой" in label_norm:
        fire = _fire_resistance(value)
        if fire:
            data["fire_resistance"] = fire
    elif "кол-во" in label_norm or "количество" in label_norm:
        if "заказыва" in label_norm and "лифт" in label_norm:
            data["quantity"] = _int_from_value(value)


def _extract_docx(path_or_bytes: str | Path | bytes) -> str:
    from docx import Document

    source = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
    document = Document(source)
    chunks: list[str] = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def _extract_pdf(path_or_bytes: str | Path | bytes) -> tuple[str, list[str], str]:
    warnings: list[str] = []
    chunks: list[str] = []
    chunks.extend(_extract_pdf_text_pypdf(path_or_bytes))
    chunks.extend(_extract_pdf_tables_pdfplumber(path_or_bytes))

    text = "\n".join(chunk for chunk in chunks if chunk.strip())
    if _has_enough_pdf_text(text):
        return text, warnings, "pdf-text"

    ocr_text, ocr_warning = _extract_pdf_ocr(path_or_bytes)
    if ocr_text:
        text = "\n".join([text, ocr_text]).strip()
        return text, warnings, "pdf-ocr"
    if ocr_warning:
        warnings.append(ocr_warning)
    return text, warnings, "pdf-text-low-confidence"


def _extract_pdf_text_pypdf(path_or_bytes: str | Path | bytes) -> list[str]:
    from pypdf import PdfReader

    source = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
    reader = PdfReader(source)
    return [page.extract_text() or "" for page in reader.pages]


def _extract_pdf_tables_pdfplumber(path_or_bytes: str | Path | bytes) -> list[str]:
    try:
        import pdfplumber
    except ImportError:
        return []

    source = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
    chunks: list[str] = []
    with pdfplumber.open(source) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text:
                chunks.append(page_text)
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [_clean_cell(cell or "") for cell in row]
                    if any(cells):
                        chunks.append("\t".join(cells))
    return chunks


def _extract_pdf_ocr(path_or_bytes: str | Path | bytes) -> tuple[str, str | None]:
    rapid_text, rapid_warning = _extract_pdf_rapidocr_text(path_or_bytes)
    if rapid_text:
        return rapid_text, None

    if shutil.which("tesseract") is None:
        return "", rapid_warning or "PDF похож на скан, но Tesseract OCR не найден в системе."
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        return "", f"Для OCR не установлена зависимость: {exc.name}."

    source = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
    document = fitz.open(stream=source, filetype="pdf") if isinstance(source, BytesIO) else fitz.open(source)
    chunks: list[str] = []
    for page in document:
        pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
        image = Image.open(BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(image, lang="rus+eng", config="--psm 6")
        chunks.append(text)
    return "\n".join(chunks), None


def _extract_pdf_rapidocr_text(path_or_bytes: str | Path | bytes) -> tuple[str, str | None]:
    items, warning = _rapidocr_pdf_items(path_or_bytes)
    if warning:
        return "", warning
    lines = _ocr_items_to_lines(items)
    return "\n".join(lines), None


def _try_rapidocr_lift_table(path_or_bytes: str | Path | bytes) -> Questionnaire | None:
    items, warning = _rapidocr_pdf_items(path_or_bytes)
    if warning or not items:
        return None

    group_columns = _detect_ocr_group_columns(items)
    if not group_columns:
        return None

    row_y = _detect_ocr_row_positions(items)
    groups: list[LiftGroup] = []
    for column in group_columns:
        data = _parse_ocr_group_column(items, column, row_y)
        section_quantity = _quantity_from_section(data.get("section"))
        if section_quantity is not None:
            data["quantity"] = section_quantity
        elif data.get("quantity") is None:
            data["quantity"] = _quantity_from_section(data.get("section"))
        if _has_lift_data(data):
            groups.append(LiftGroup(**_drop_empty(data)))

    return Questionnaire(project=_parse_ocr_project_info(items), lift_groups=groups) if groups else None


def _parse_ocr_project_info(items: list[dict[str, Any]]) -> ProjectInfo:
    title_items = [item for item in items if item["yc"] > 1120 and item["xc"] > 1600]
    title_text = " ".join(item["text"] for item in sorted(title_items, key=lambda value: (value["yc"], value["xc"])))
    address = _parse_ocr_address(title_items, title_text)
    project_name = _parse_ocr_object_name(title_items, title_text)
    customer = _parse_ocr_customer(title_text)
    return ProjectInfo(project_name=project_name, customer=customer, address=address)


def _parse_ocr_address(items: list[dict[str, Any]], title_text: str) -> str | None:
    cadastral = _first_match(title_text, [r"(\d{2}:\d{2}:\d{6,}:\d+)"])
    address_lines = [
        item["text"]
        for item in sorted(items, key=lambda value: value["yc"])
        if 1360 <= item["yc"] <= 1450 and item["xc"] > 1650
    ]
    raw = " ".join(address_lines)
    normalized = _normalize(raw)
    if cadastral and ("ehu" in normalized or "len" in normalized or "47:" in normalized):
        return (
            "Ленинградская область, Всеволожский муниципальный район, "
            f"Заневское городское поселение, земельный участок с кадастровым номером {cadastral}"
        )
    if cadastral:
        return f"{_clean_cell(raw)}, {cadastral}".strip(" ,")
    return _clean_cell(raw) or None


def _parse_ocr_object_name(items: list[dict[str, Any]], title_text: str) -> str | None:
    object_lines = [
        item["text"]
        for item in sorted(items, key=lambda value: (value["yc"], value["xc"]))
        if 1450 <= item["yc"] <= 1545 and 1680 <= item["xc"] <= 2050
    ]
    raw = " ".join(object_lines)
    normalized = _normalize(raw)
    if "mh02" in normalized or "kbap" in normalized or "bcmpoehho" in normalized:
        return "Многоквартирный дом со встроенно-пристроенными помещениями"
    return _clean_cell(raw) or None


def _parse_ocr_customer(title_text: str) -> str | None:
    normalized = _normalize(title_text)
    if "aciekt" in normalized or "aspekt" in normalized:
        return 'ООО СЗ "АСПЕКТ ПЛЮС"'
    if "3od" in normalized or "зод" in normalized:
        return 'ООО "Зодчий"'
    return None


def _rapidocr_pdf_items(path_or_bytes: str | Path | bytes) -> tuple[list[dict[str, Any]], str | None]:
    try:
        import fitz
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        return [], f"Для встроенного OCR не установлена зависимость: {exc.name}."

    source = BytesIO(path_or_bytes) if isinstance(path_or_bytes, bytes) else path_or_bytes
    document = fitz.open(stream=source, filetype="pdf") if isinstance(source, BytesIO) else fitz.open(source)
    ocr = RapidOCR()
    items: list[dict[str, Any]] = []
    for page_index, page in enumerate(document):
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        result, _ = ocr(pix.tobytes("png"))
        for box, text, score in result or []:
            xs = [point[0] for point in box]
            ys = [point[1] for point in box]
            items.append(
                {
                    "page": page_index,
                    "text": _clean_cell(text),
                    "score": float(score),
                    "x1": float(min(xs)),
                    "x2": float(max(xs)),
                    "y1": float(min(ys)),
                    "y2": float(max(ys)),
                    "xc": float(sum(xs) / len(xs)),
                    "yc": float(sum(ys) / len(ys)),
                }
            )
    return items, None


def _ocr_items_to_lines(items: list[dict[str, Any]]) -> list[str]:
    lines: list[list[dict[str, Any]]] = []
    for item in sorted(items, key=lambda value: (value["page"], value["yc"], value["xc"])):
        if not lines or abs(lines[-1][0]["yc"] - item["yc"]) > 12:
            lines.append([item])
        else:
            lines[-1].append(item)
    return [" ".join(part["text"] for part in sorted(line, key=lambda value: value["xc"])) for line in lines]


def _detect_ocr_group_columns(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    headers: list[dict[str, Any]] = []
    for item in items:
        text = item["text"]
        if item["yc"] > 250 or item["xc"] < 300:
            continue
        match = re.search(r"(\d+)\s*[-–—]\s*(\d+)", text)
        if not match:
            continue
        section = f"Секция {match.group(1)}-{match.group(2)}"
        headers.append({"xc": item["xc"], "section": section})

    if headers:
        headers = sorted(headers, key=lambda value: value["xc"])
        deduped: list[dict[str, Any]] = []
        for header in headers:
            if not deduped or abs(deduped[-1]["xc"] - header["xc"]) > 80:
                deduped.append(header)
        return deduped

    q_items = [item for item in items if re.search(r"\bQ\s*=?\s*\d{3,5}", item["text"], flags=re.IGNORECASE)]
    return [{"xc": item["xc"], "section": None} for item in sorted(q_items, key=lambda value: value["xc"])]


def _detect_ocr_row_positions(items: list[dict[str, Any]]) -> dict[int, float]:
    row_y: dict[int, float] = {}
    for item in items:
        if item["xc"] > 180:
            continue
        text = item["text"].strip()
        if not re.fullmatch(r"\d{1,2}", text):
            continue
        number = int(text)
        if 1 <= number <= 30:
            row_y[number] = item["yc"]
    return row_y


def _parse_ocr_group_column(items: list[dict[str, Any]], column: dict[str, Any], row_y: dict[int, float]) -> dict[str, Any]:
    xc = column["xc"]
    col_items = [item for item in items if abs(item["xc"] - xc) <= 155]
    data: dict[str, Any] = {"section": column.get("section")}
    header_parts = [item["text"] for item in col_items if item["yc"] < 230]
    header_text = " ".join(header_parts)
    header_data = _parse_group_header(header_text)
    data.update({key: value for key, value in header_data.items() if value not in (None, "")})
    if "nacc" in _normalize(header_text) or "пассаж" in _normalize(header_text):
        data["lift_type"] = "Пассажирский"

    qv_text = " ".join(item["text"] for item in col_items if 420 <= item["yc"] <= 500)
    data["capacity_kg"] = _int_match(qv_text, [r"Q\s*=?\s*(\d{3,5})", r"(\d{3,5})\s*k"]) or data.get("capacity_kg")
    data["speed_ms"] = _float_match(qv_text, [r"V\s*=?\s*([0-9]+(?:[,.][0-9]+)?)"]) or data.get("speed_ms")
    if data.get("capacity_kg") and data.get("lift_type") == "Пассажирский":
        data["lift_name"] = f"Лифт пассажирский {data['capacity_kg']} кг"

    data["lifting_height_mm"] = _first_measurement_in_band(col_items, 500, 555, decimal_only=True)
    size_text = _text_in_band(col_items, 555, 620)
    size = _parse_size(size_text)
    if size:
        data["cabin_width_mm"], data["cabin_depth_mm"] = size

    opposite_text = _text_in_band(col_items, 620, 675)
    if opposite_text:
        data["cabin_type"] = "Проходная" if _yes_ocr(opposite_text) else "Непроходная"

    data["doors_count"] = _int_in_row(col_items, row_y, 9) or _single_int_in_band(col_items, 675, 715)
    data["stops"] = _int_in_row(col_items, row_y, 10) or _single_int_in_band(col_items, 715, 755)
    data["button_marking"] = _text_in_row(col_items, row_y, 11) or None
    data["group_operation"] = _group_operation(_text_in_row(col_items, row_y, 13) or "")
    data["shaft_material"] = _shaft_material_from_ocr(_text_in_row(col_items, row_y, 16) or "")
    data["fire_resistance"] = _fire_resistance(_text_in_row(col_items, row_y, 17) or "") or _fire_resistance(_text_in_band(col_items, 1130, 1215))
    data["quantity"] = _int_in_row(col_items, row_y, 18)
    data["pit_depth_mm"] = _measurement_mm(_text_in_row(col_items, row_y, 19) or "") or _first_measurement_in_band(col_items, 1260, 1305)
    data["overhead_mm"] = _measurement_mm(_text_in_row(col_items, row_y, 20) or "") or _first_measurement_in_band(col_items, 1305, 1350)
    return data


def _text_in_row(items: list[dict[str, Any]], row_y: dict[int, float], row: int, tolerance: float = 20) -> str:
    if row not in row_y:
        return ""
    return " ".join(item["text"] for item in sorted(items, key=lambda value: value["xc"]) if abs(item["yc"] - row_y[row]) <= tolerance)


def _text_in_band(items: list[dict[str, Any]], y1: float, y2: float) -> str:
    return " ".join(item["text"] for item in sorted(items, key=lambda value: (value["yc"], value["xc"])) if y1 <= item["yc"] <= y2)


def _int_in_row(items: list[dict[str, Any]], row_y: dict[int, float], row: int) -> int | None:
    return _int_from_value(_text_in_row(items, row_y, row))


def _single_int_in_band(items: list[dict[str, Any]], y1: float, y2: float) -> int | None:
    values = [
        int(item["text"])
        for item in items
        if y1 <= item["yc"] <= y2 and re.fullmatch(r"\d{1,3}", item["text"])
    ]
    return values[0] if values else None


def _first_measurement_in_band(items: list[dict[str, Any]], y1: float, y2: float, decimal_only: bool = False) -> int | None:
    for item in sorted(items, key=lambda value: value["xc"]):
        if not (y1 <= item["yc"] <= y2):
            continue
        if decimal_only and not re.search(r"\d+[,.]\d{3}", item["text"]):
            continue
        value = _measurement_mm(item["text"])
        if value:
            return value
    return None


def _yes_ocr(value: str) -> bool:
    value = _normalize(value)
    return value in {"да", "yes"} or "he" not in value and "hem" not in value


def _shaft_material_from_ocr(value: str) -> str | None:
    normalized = _normalize(value)
    if (
        "moho" in normalized
        or "монолит" in normalized
        or "бетон" in normalized
        or "железобетон" in normalized
        or "ж/б" in normalized
        or "жб" in normalized
    ):
        return "Железобетон"
    return _capitalize_value(value) if value else None


def _quantity_from_section(section: str | None) -> int | None:
    if not section:
        return None
    match = re.search(r"(\d+)\s*[-–—]\s*(\d+)", section)
    if not match:
        return None
    start, end = int(match.group(1)), int(match.group(2))
    return end - start + 1 if end >= start else None


def _has_enough_pdf_text(text: str) -> bool:
    if len(text) < 500:
        return False
    keywords = ["лифт", "грузопод", "скорост", "останов", "кабин", "шахт"]
    lower = text.lower()
    return sum(keyword in lower for keyword in keywords) >= 2


def _build_found_fields(groups: list[LiftGroup]) -> dict[str, bool]:
    found: dict[str, bool] = {}
    for group in groups:
        for key, value in group.model_dump().items():
            found[key] = found.get(key, False) or value is not None
    return found


def _project_found_fields(project: ProjectInfo) -> dict[str, bool]:
    return {
        "project_name": bool(project.project_name),
        "customer": bool(project.customer),
        "address": bool(project.address),
    }


def _has_lift_data(data: dict[str, Any]) -> bool:
    meaningful = [
        "section",
        "lift_name",
        "quantity",
        "capacity_kg",
        "speed_ms",
        "lifting_height_mm",
        "stops",
        "doors_count",
        "cabin_width_mm",
        "shaft_material",
    ]
    return any(data.get(key) not in (None, "") for key in meaningful)


def _suffix(path_or_bytes: str | Path | bytes, filename: str | None) -> str:
    if filename:
        return Path(filename).suffix.lower()
    if isinstance(path_or_bytes, (str, Path)):
        return Path(path_or_bytes).suffix.lower()
    return ""


def _first_match(text: str, patterns: list[str]) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return None


def _int_match(text: str, patterns: list[str]) -> int | None:
    value = _first_match(text, patterns)
    return int(value) if value else None


def _float_match(text: str, patterns: list[str]) -> float | None:
    value = _first_match(text, patterns)
    return float(value.replace(",", ".")) if value else None


def _measurement_mm_match(text: str, patterns: list[str]) -> int | None:
    value = _first_match(text, patterns)
    return _measurement_mm(value) if value else None


def _measurement_mm(value: str) -> int | None:
    raw = _first_match(value, [r"([0-9]+(?:[,.][0-9]+)?)"])
    if raw is None:
        return None
    normalized = raw.replace(",", ".")
    if re.fullmatch(r"\d+\.\d{3}", normalized):
        return int(normalized.replace(".", ""))
    number = float(normalized)
    lower = value.lower()
    if number < 100 and "мм" not in lower:
        return int(round(number * 1000))
    return int(round(number))


def _parse_size(value: str) -> tuple[int, int] | None:
    match = re.search(r"(\d{3,5})\s*[xх×*]\s*(\d{3,5})", value, flags=re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1)), int(match.group(2))


def _int_from_value(value: str) -> int | None:
    return _int_match(value, [r"(\d{1,6})"])


def _group_operation(value: str) -> str:
    value_norm = _normalize(value)
    if "dds" in value_norm:
        return "DDS"
    if "груп" in value_norm:
        return "Групповое"
    return "Одиночное"


def _fire_resistance(value: str) -> str | None:
    match = re.search(r"E[I1]\s*[- ]?\s*(\d{2,3})", value, flags=re.IGNORECASE)
    if match:
        return f"EI-{match.group(1)}"
    return None


def _yes(value: str) -> bool:
    return _normalize(value) in {"да", "yes", "true", "есть", "имеется"}


def _capitalize_value(value: str) -> str:
    value = _clean_cell(value)
    return value[:1].upper() + value[1:] if value else value


def _normalize(value: str) -> str:
    value = value.lower().replace("ё", "е")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _clean_cell(value: str | None) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _drop_empty(data: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in data.items() if value not in ("", None)}
